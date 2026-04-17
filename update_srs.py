import docx
import copy
from docx.table import Table
from docx.shared import Inches

def insert_paragraph_before(reference_paragraph, text, style=None):
    new_p = docx.oxml.OxmlElement('w:p')
    reference_paragraph._p.addprevious(new_p)
    p = docx.text.paragraph.Paragraph(new_p, reference_paragraph._parent)
    if text:
        p.add_run(text)
    if style is not None:
        p.style = style
    return p

def insert_table_before(reference_paragraph, template_table, fields):
    new_tbl = copy.deepcopy(template_table._tbl)
    reference_paragraph._p.addprevious(new_tbl)
    table_obj = Table(new_tbl, reference_paragraph._parent)
    
    # Optional: adjust row count or just replace Text
    # We will just write fields sequentially into the table starting from row 1
    # Header is row 0
    # ensure we have enough rows
    while len(table_obj.rows) <= len(fields):
        new_row = copy.deepcopy(table_obj.rows[-1]._tr)
        table_obj._tbl.append(new_row)
    
    # clear all rows from 1 onwards
    for i in range(1, len(table_obj.rows)):
        table_obj.cell(i, 0).text = ""
        table_obj.cell(i, 1).text = ""

    # populate
    for i, (fname, fdesc) in enumerate(fields):
        try:
            table_obj.cell(i+1, 0).text = fname
            table_obj.cell(i+1, 1).text = fdesc
        except:
            pass

data = [
    {
        "feature": "3.1 Parents Portal",
        "subfeatures": [
            {
                "name": "3.1.1 Authentication",
                "screens": [
                    {
                        "name": "3.1.1.1 Parent Landing & Login",
                        "png": "Document/03_Design/03_UI/UX/parent_landing_login.png",
                        "desc": "This screen allows the Parents to view the landing page, log in using their credentials (e.g., phone number), or proceed to the self-service booking portal as a guest.",
                        "fields": [
                            ("(1)", "Phone Number: Enter registered phone number."),
                            ("(2)", "OTP/Password: Enter the OTP sent to the phone.")
                        ]
                    }
                ]
            },
            {
                "name": "3.1.2 Self-Service Booking",
                "screens": [
                    {
                        "name": "3.1.2.1 Booking Portal",
                        "png": "Document/03_Design/03_UI/UX/parent_booking_portal.png",
                        "desc": "This screen allows the Parents to book appointments online by selecting a doctor, choosing an available time slot, and providing patient demographics (Name, DOB, Symptoms). Features self-service capabilities without receptionist intervention.",
                        "fields": [
                            ("(1)", "Doctor Selection: Choose from the list of available doctors."),
                            ("(2)", "Time Slot: Available times for the selected doctor."),
                            ("(3)", "Patient Information: Name, DOB, and Reason for Visit.")
                        ]
                    },
                    {
                        "name": "3.1.2.2 Appointment Confirmed",
                        "png": "Document/03_Design/03_UI/UX/parent_appointment_confirmed.png",
                        "desc": "This screen allows the Parents to receive their appointment confirmation details after booking. It displays the appointment time, doctor, and an option to return to the dashboard.",
                        "fields": [
                            ("(1)", "Status Message: Confirmation status (e.g., Confirmed)."),
                            ("(2)", "Appointment Details: Reference ID, Date, Time, Doctor Name.")
                        ]
                    }
                ]
            },
            {
                "name": "3.1.3 Patient EMR Portal",
                "screens": [
                    {
                        "name": "3.1.3.1 Parent Dashboard",
                        "png": "Document/03_Design/03_UI/UX/parent_dashboard.png",
                        "desc": "This screen allows the Parents to view an overview of upcoming appointments and access past medical records for their children.",
                        "fields": [
                            ("(1)", "Upcoming Appointments: List of future confirmed visits."),
                            ("(2)", "Recent Visits: Quick links to recent clinical notes.")
                        ]
                    },
                    {
                        "name": "3.1.3.2 Clinical Notes & Prescriptions",
                        "png": "Document/03_Design/03_UI/UX/parent_clinical_notes.png",
                        "desc": "This screen allows the Parents to view clinical notes and prescriptions issued by the doctor after a visit. Promotes 100% EMR traceability and frictionless access.",
                        "fields": [
                            ("(1)", "Clinical Notes: Doctor's observations and diagnosis."),
                            ("(2)", "Prescription: Prescribed medications, quantity, and dosage instructions.")
                        ]
                    }
                ]
            }
        ]
    },
    {
        "feature": "3.2 Clinic Staff Operations",
        "subfeatures": [
            {
                "name": "3.2.1 Staff Authentication",
                "screens": [
                    {
                        "name": "3.2.1.1 Staff Login",
                        "png": "Document/03_Design/03_UI/UX/shared_staff_login.png",
                        "desc": "This screen allows clinic staff (Receptionist, Doctor) to securely log into the HCMS system using authenticated credentials before accessing their respective role-based dashboards.",
                        "fields": [
                            ("(1)", "Username: Enter staff account username."),
                            ("(2)", "Password: Enter secure password."),
                            ("(3)", "Role: Auto-detected or selected after login.")
                        ]
                    }
                ]
            }
        ]
    },
    {
        "feature": "3.3 Receptionist Portal",
        "subfeatures": [
            {
                "name": "3.3.1 Appointment Management",
                "screens": [
                    {
                        "name": "3.3.1.1 Appointment Dashboard",
                        "png": "Document/03_Design/03_UI/UX/receptionist_appointment_dashboard.png",
                        "desc": "This screen allows the Receptionist to view the daily appointment pipeline, monitor statuses (Pending, Confirmed), and manually confirm online bookings.",
                        "fields": [
                            ("(1)", "Appointment List: Displays all bookings for the day."),
                            ("(2)", "Status Filter: View Pending, Confirmed, Walk-in, etc."),
                            ("(3)", "Confirm Action: Button to manually confirm pending bookings.")
                        ]
                    },
                    {
                        "name": "3.3.1.2 Register Walk-in Appointment",
                        "png": "Document/03_Design/03_UI/UX/receptionist_register_walkin.png",
                        "desc": "This screen allows the Receptionist to rapidly register an appointment for walk-in patients by entering the required demographic information and assigning them to an available doctor queue.",
                        "fields": [
                            ("(1)", "Patient Information: Name, DOB, Contact details."),
                            ("(2)", "Assigned Doctor: Staff member assigned to the walk-in case.")
                        ]
                    }
                ]
            },
            {
                "name": "3.3.2 Billing & Payment",
                "screens": [
                    {
                        "name": "3.3.2.1 Billing Invoice",
                        "png": "Document/03_Design/03_UI/UX/receptionist_billing_invoice.png",
                        "desc": "This screen allows the Receptionist to view generated invoices post-consultation and systematically update the payment status using a optimized '1-Click Paid' action.",
                        "fields": [
                            ("(1)", "Invoice ID & Patient Name: Identifier for the bill."),
                            ("(2)", "Amount Due: Calculated from consultation fee and medicines."),
                            ("(3)", "1-Click Pay: Button to instantly update status from Unpaid to Paid.")
                        ]
                    }
                ]
            }
        ]
    },
    {
        "feature": "3.4 Doctor Portal",
        "subfeatures": [
            {
                "name": "3.4.1 Clinical Dashboard",
                "screens": [
                    {
                        "name": "3.4.1.1 Patient Queue Dashboard",
                        "png": "Document/03_Design/03_UI/UX/doctor_queue_dashboard.png",
                        "desc": "This screen allows the Doctor to view the real-time queue of waiting patients and select the next patient for consultation.",
                        "fields": [
                            ("(1)", "Patient Queue: Ordered list of waiting patients."),
                            ("(2)", "Wait Time: Track time elapsed since check-in.")
                        ]
                    }
                ]
            },
            {
                "name": "3.4.2 EMR Management",
                "screens": [
                    {
                        "name": "3.4.2.1 Clinical EMR & Consultation",
                        "png": "Document/03_Design/03_UI/UX/doctor_clinical_emr.png",
                        "desc": "This screen allows the Doctor to seamlessly access patient medical history, record detailed notes during clinical consultations, and immediately issue electronic prescriptions utilizing autocomplete fields to minimize typing.",
                        "fields": [
                            ("(1)", "Medical History: Previous clinical records."),
                            ("(2)", "Diagnosis (Autocomplete): Fast-entry field for recording diagnosis."),
                            ("(3)", "Prescription: Dynamic table to add medications.")
                        ]
                    },
                    {
                        "name": "3.4.2.2 Manage Patient Profile",
                        "png": "Document/03_Design/03_UI/UX/doctor_manage_profile.png",
                        "desc": "This screen allows the Doctor (or authorized staff) to update master records and demographic information for specific patients within the EMR.",
                        "fields": [
                            ("(1)", "Demographics: Basic info like age, weight, gender."),
                            ("(2)", "Chronic Conditions: Any pre-existing health issues."),
                            ("(3)", "Allergies: Drug or food allergies.")
                        ]
                    }
                ]
            }
        ]
    }
]

doc = docx.Document("Document/02_Requirements/HCMS_SRS_v0.1.docx")
p_target = doc.paragraphs[120] # "3.2 User Authentication" -> will insert before this.
t_template = doc.tables[20]

# Pre-clear paragraphs 109 to 119
for i in range(109, 120):
    doc.paragraphs[i].text = ""

# Now build the structure
for feature in data:
    insert_paragraph_before(p_target, feature['feature'], style='Heading 3')
    for sub in feature['subfeatures']:
        insert_paragraph_before(p_target, sub['name'], style='Heading 4')
        for screen in sub['screens']:
            insert_paragraph_before(p_target, screen['name'], style='Heading 5')
            
            # Picture
            pic_p = insert_paragraph_before(p_target, "", style='Normal')
            run = pic_p.add_run()
            run.add_picture(screen['png'], width=Inches(6.0))
            
            # Description
            insert_paragraph_before(p_target, screen['desc'], style='Normal')
            insert_paragraph_before(p_target, "Field Description", style='Normal')
            
            # Table
            insert_table_before(p_target, t_template, screen['fields'])

doc.save("Document/02_Requirements/HCMS_SRS_v0.1.docx")
print("Document updated.")
