import docx
from docx.shared import Pt

file_path = r'e:\Documents\Ky_8\Co Dieu AISDLC\HCMS---Healthcare-Clinic-System\Document\02_Requirements\HCMS_SRS_v0.1.docx'
doc = docx.Document(file_path)

ucs_data = [
    {
        "feature": "2.1 Self-Service Booking Portal",
        "ucs": [
            {
                "title": "2.1.1 UC-01 Book Appointment Online",
                "primary_actor": "Parents", "secondary_actor": "HCMS System",
                "desc": "As a parent, I want to view available doctors and time slots to book an appointment online, reducing waiting time.",
                "pre": "Parent accesses the Booking Portal.",
                "post": "APPOINTMENT is created with status 'Pending'.",
                "norm": "1. Parent selects preferred Doctor and time slot.\n2. Parent confirms the selection.\n3. System proceeds to demographic form (UC-02).",
                "alt": "None"
            },
            {
                "title": "2.1.2 UC-02 Provide Patient Demographics",
                "primary_actor": "Parents", "secondary_actor": "HCMS System",
                "desc": "As a parent, I want to enter my child's information so the clinic can properly identify the patient.",
                "pre": "Parent has selected a time slot in UC-01.",
                "post": "PATIENT entity is updated/created.",
                "norm": "1. System requests patient name, DOB, and parent contact.\n2. Parent inputs details and submits.\n3. System validates and saves PATIENT and APPOINTMENT data.",
                "alt": "Step 3_Validation Error:\nSystem highlights missing fields."
            }
        ]
    },
    {
        "feature": "2.2 Appointment Notification",
        "ucs": [
            {
                "title": "2.2.1 UC-03 Receive Appointment Confirmation",
                "primary_actor": "Parents", "secondary_actor": "HCMS System",
                "desc": "As a parent, I want to receive confirmation once the receptionist verifies my booking.",
                "pre": "Receptionist confirmed the APPOINTMENT.",
                "post": "Parent receives a notification.",
                "norm": "1. Receptionist changes status to Confirmed.\n2. System sends notification to Parent.\n3. Parent views confirmation.",
                "alt": "None"
            }
        ]
    },
    {
        "feature": "2.3 Patient EMR Portal",
        "ucs": [
            {
                "title": "2.3.1 UC-04 View Clinical Notes & Prescriptions",
                "primary_actor": "Parents", "secondary_actor": "HCMS System",
                "desc": "As a parent, I want to view my child's consultation notes and medication dosage on my phone.",
                "pre": "Consultation finalized (VISIT and PRESCRIPTION exist).",
                "post": "None",
                "norm": "1. Parent accesses the EMR Portal via link.\n2. System displays the latest VISIT notes and PRESCRIPTION.\n3. Parent reads the instructions.",
                "alt": "None"
            }
        ]
    },
    {
        "feature": "2.4 Appointment Management",
        "ucs": [
            {
                "title": "2.4.1 UC-05 View Appointment Dashboard",
                "primary_actor": "Receptionist", "secondary_actor": "HCMS System",
                "desc": "As a receptionist, I want to see an overview of all appointments for the day.",
                "pre": "Receptionist logged in.",
                "post": "None",
                "norm": "1. Receptionist navigates to Dashboard.\n2. System queries and displays APPOINTMENT queue list.",
                "alt": "None"
            },
            {
                "title": "2.4.2 UC-06 Confirm Appointment",
                "primary_actor": "Receptionist", "secondary_actor": "HCMS System",
                "desc": "As a receptionist, I want to confirm pending online bookings to ensure schedule accuracy.",
                "pre": "Dashboard displays a 'Pending' APPOINTMENT.",
                "post": "APPOINTMENT status becomes 'Confirmed'.",
                "norm": "1. Receptionist selects a pending appointment.\n2. Receptionist clicks Confirm.\n3. System updates APPOINTMENT status.",
                "alt": "None"
            },
            {
                "title": "2.4.3 UC-07 Register Walk-in Appointment",
                "primary_actor": "Receptionist", "secondary_actor": "HCMS System",
                "desc": "As a receptionist, I want to register walk-in patients quickly at the desk.",
                "pre": "Parent arrives at the clinic without booking.",
                "post": "APPOINTMENT and PATIENT records are created.",
                "norm": "1. Receptionist clicks 'New Walk-in'.\n2. Receptionist enters patient details and assigns to Doctor queue.\n3. System saves entities and updates Dashboard.",
                "alt": "None"
            }
        ]
    },
    {
        "feature": "2.5 Billing & Payment",
        "ucs": [
            {
                "title": "2.5.1 UC-08 View Billing Invoice",
                "primary_actor": "Receptionist", "secondary_actor": "HCMS System",
                "desc": "As a receptionist, I want to view the auto-generated checkout invoice for a patient.",
                "pre": "Doctor generated the BILLING record after visit.",
                "post": "None",
                "norm": "1. Receptionist selects the 'Unpaid' bill.\n2. System displays cost breakdown.\n3. Receptionist confirms amount with Parent.",
                "alt": "None"
            },
            {
                "title": "2.5.2 UC-09 Update Payment Status (1-Click Paid)",
                "primary_actor": "Receptionist", "secondary_actor": "Parents",
                "desc": "As a receptionist, I want to mark the bill as paid with 1-click.",
                "pre": "Parent completes payment.",
                "post": "BILLING status is 'Paid'.",
                "norm": "1. Receptionist clicks the 'Paid' button.\n2. System updates the BILLING record.\n3. System logs revenue.",
                "alt": "Step 1_User payment fails:\nReceptionist leaves as 'Unpaid'."
            }
        ]
    },
    {
        "feature": "2.6 Clinical Dashboard",
        "ucs": [
            {
                "title": "2.6.1 UC-10 View Patient Queue",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to view the queue of waiting patients assigned to me.",
                "pre": "Doctor logged in.",
                "post": "None",
                "norm": "1. Doctor accesses Clinical Dashboard.\n2. System displays list of waiting APPOINTMENTS.\n3. Doctor selects the next patient.",
                "alt": "None"
            }
        ]
    },
    {
        "feature": "2.7 EMR Access",
        "ucs": [
            {
                "title": "2.7.1 UC-11 Access Patient Medical History",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to review past visits and diagnoses to make informed clinical decisions.",
                "pre": "Doctor selects a patient from the queue.",
                "post": "None",
                "norm": "1. Doctor clicks on Patient Profile.\n2. System queries past VISITs and PRESCRIPTIONs.\n3. System displays historical summary.",
                "alt": "None"
            }
        ]
    },
    {
        "feature": "2.8 EMR \u2013 Visit & Prescription Management",
        "ucs": [
            {
                "title": "2.8.1 UC-12 Record Clinical Consultation",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to type clinical notes and diagnoses directly into the EMR.",
                "pre": "Doctor accesses Patient EMR.",
                "post": "VISIT entity is created/updated.",
                "norm": "1. Doctor inputs symptoms, notes, and diagnosis.\n2. Doctor clicks Save.\n3. System safely stores the VISIT record.",
                "alt": "None"
            },
            {
                "title": "2.8.2 UC-13 Issue Electronic Prescription",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to quickly add medications using autocomplete.",
                "pre": "Doctor is in the active VISIT session.",
                "post": "PRESCRIPTION entity is created.",
                "norm": "1. Doctor types few letters of medication.\n2. System suggests drug names (Autocomplete).\n3. Doctor selects drug, specifies dosage, clicks Finalize.\n4. System safely stores the PRESCRIPTION record.",
                "alt": "None"
            }
        ]
    },
    {
        "feature": "2.9 EMR \u2013 Patient Management",
        "ucs": [
            {
                "title": "2.9.1 UC-14 Manage Patient Profile",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to correct or update the core demographic/allergy information of the patient.",
                "pre": "Doctor accesses Patient EMR.",
                "post": "PATIENT entity is updated.",
                "norm": "1. Doctor clicks Edit Profile.\n2. Doctor updates allergy/baseline info.\n3. System saves the PATIENT record.",
                "alt": "None"
            }
        ]
    }
]

def format_cell(cell, text, bold=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.bold = bold

def create_table_element(uc):
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    format_cell(table.cell(0, 0), 'Primary Actors')
    format_cell(table.cell(0, 1), uc['primary_actor'])
    format_cell(table.cell(0, 2), 'Secondary Actors')
    format_cell(table.cell(0, 3), uc['secondary_actor'])
    
    format_cell(table.cell(1, 0), 'Description')
    table.cell(1, 1).merge(table.cell(1, 3))
    format_cell(table.cell(1, 1), uc['desc'])
    
    format_cell(table.cell(2, 0), 'Preconditions')
    table.cell(2, 1).merge(table.cell(2, 3))
    format_cell(table.cell(2, 1), uc['pre'])
    
    format_cell(table.cell(3, 0), 'Postconditions')
    table.cell(3, 1).merge(table.cell(3, 3))
    format_cell(table.cell(3, 1), uc['post'])
    
    format_cell(table.cell(4, 0), 'Normal Sequence/Flow')
    table.cell(4, 1).merge(table.cell(4, 3))
    format_cell(table.cell(4, 1), uc['norm'])
    
    format_cell(table.cell(5, 0), 'Alternative Sequences/Flows')
    table.cell(5, 1).merge(table.cell(5, 3))
    format_cell(table.cell(5, 1), uc['alt'])
    
    return table._tbl

body = doc._body._body

start_elem = None
end_elem = None

for p in doc.paragraphs:
    txt = p.text.strip()
    # Identifying bounds
    if txt.startswith("2.1 <<Feature Name1>>"):
        start_elem = p._element
    elif txt.startswith("2.2 Xyz Feature"):
        end_elem = p._element

if start_elem is not None and end_elem is not None:
    # Gather elements to delete, from start_elem inclusive up to end_elem exclusive
    elements_to_delete = []
    in_range = False
    for elem in body:
        if elem == start_elem:
            in_range = True
        if elem == end_elem:
            in_range = False
        if in_range:
            elements_to_delete.append(elem)
            
    # Delete them
    for elem in elements_to_delete:
        body.remove(elem)
        
    # Now end_elem is our anchor. We insert BEFORE end_elem.
    for feature_obj in ucs_data:
        # Create Feature Title
        feat_p = doc.add_paragraph(feature_obj["feature"])
        feat_p.style = doc.styles['Heading 2']
        for r in feat_p.runs: r.font.name = 'Arial'
        end_elem.addprevious(feat_p._p)
        
        for uc in feature_obj["ucs"]:
            # UC Title
            uc_p = doc.add_paragraph(uc["title"])
            uc_p.style = doc.styles['Heading 3']
            for r in uc_p.runs: 
                r.font.name = 'Arial'
                r.bold = True
            end_elem.addprevious(uc_p._p)
            
            # UC Table
            tbl_xml = create_table_element(uc)
            end_elem.addprevious(tbl_xml)
            
            # Spacer
            spacer = doc.add_paragraph("")
            end_elem.addprevious(spacer._p)

    doc.save(file_path)
    print("Done! Replaced 14 use cases.")
else:
    print("Error: Could not find markers.")
