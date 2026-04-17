import docx
from docx.shared import Pt
import json

file_path = r'e:\Documents\Ky_8\Co Dieu AISDLC\HCMS---Healthcare-Clinic-System\Document\02_Requirements\HCMS_SRS_v0.1.docx'
doc = docx.Document(file_path)

ucs_data = [
    {
        "feature": "2.1 Self-Service Booking Portal",
        "ucs": [
            {
                "title": "2.1.1 UC-01 Book Appointment Online",
                "primary_actor": "Parents", "secondary_actor": "HCMS System",
                "desc": "As a parent, I want to view available doctors and time slots to book an appointment online so that I can save time and avoid waiting to call the clinic.",
                "pre": "Parent accesses the Self-Service Booking Portal.",
                "post": "●\tAPPOINTMENT is tentatively held.\n●\tSystem proceeds to demographics screen.",
                "norm": "Book Appointment\n1.\tParent accesses the Self-Service Booking Portal via web or mobile browser.\n2.\tSystem displays available Doctors and time slots for the upcoming 7 days.\n3.\tParent selects a preferred Doctor and an available time slot.\n4.\tParent clicks the 'Continue' button.\n5.\tSystem temporarily holds the selected time slot.\n6.\tSystem directs Parent to the Patient Demographic screen (UC-02).",
                "alt": "Step 4_Time Slot Unavailable\n1.\tSystem checks slot availability. If it was just booked by another user, System displays an error message (MSG01).\n2.\tSystem refreshes available time slots.\n3.\tReturn to step 2 of normal flow."
            },
            {
                "title": "2.1.2 UC-02 Provide Patient Demographics",
                "primary_actor": "Parents", "secondary_actor": "HCMS System",
                "desc": "As a parent, I want to enter my child's information so the clinic can properly identify the patient and confirm the booking.",
                "pre": "Parent has successfully selected a time slot (UC-01).",
                "post": "●\tPATIENT entity is updated/created.\n●\tAPPOINTMENT entity is created with status 'Pending'.",
                "norm": "Provide Demographics\n1.\tSystem displays the Patient Demographic form.\n2.\tParent inputs the child's details (Full Name, Date of Birth, Gender) and parent contact (Phone Number, Email), along with Reason for Visit.\n3.\tParent clicks the 'Submit Booking' button.\n4.\tSystem validates the inputted information (BR-01, BR-02).\n5.\tSystem creates/updates the PATIENT record in the database.\n6.\tSystem creates a new APPOINTMENT with status 'Pending' and links to the PATIENT.\n7.\tSystem updates the Receptionist's Dashboard with the new APPOINTMENT.\n8.\tSystem directs Parent to the Booking Success screen.",
                "alt": "Step 4_Validation Errors\n1.\tParent leaves mandatory fields blank (MSG02) or inputs invalid phone/email format (MSG03).\n2.\tSystem highlights invalid fields with red text and displays error messages.\n3.\tReturn to step 2 of normal flow."
            }
        ]
    },
    {
        "feature": "2.2 Appointment Notification",
        "ucs": [
            {
                "title": "2.2.1 UC-03 Receive Appointment Confirmation",
                "primary_actor": "Parents", "secondary_actor": "Receptionist",
                "desc": "As a parent, I want to receive confirmation once the receptionist verifies my booking so I know my appointment is secured.",
                "pre": "Receptionist confirmed the 'Pending' APPOINTMENT.",
                "post": "●\tParent receives a notification message.\n●\tActivity Log tracks the sent message.",
                "norm": "Receive Confirmation\n1.\tReceptionist changes the APPOINTMENT status to 'Confirmed'.\n2.\tSystem generates an SMS/Email notification containing appointment details (Doctor, Time, Clinic Address).\n3.\tSystem sends the notification to the Parent's registered contact info.\n4.\tParent receives and views the confirmation message on their device.",
                "alt": "Step 1.1_Appointment Rejected\n1.\tReceptionist changes the APPOINTMENT status to 'Cancelled' (e.g., doctor unavailable).\n2.\tSystem sends a cancellation notification with the corresponding reason to the Parent."
            }
        ]
    },
    {
        "feature": "2.3 Patient EMR Portal",
        "ucs": [
            {
                "title": "2.3.1 UC-04 View Clinical Notes & Prescriptions",
                "primary_actor": "Parents", "secondary_actor": "HCMS System",
                "desc": "As a parent, I want to view my child's consultation notes and medication dosage on my phone, to easily follow the doctor's instructions.",
                "pre": "Consultation finalized (VISIT and PRESCRIPTION records exist).",
                "post": "●\tParent successfully authenticates via OTP and views records.",
                "norm": "View Patient EMR\n1.\tParent accesses the EMR Portal via the secure link sent to their phone after the visit.\n2.\tSystem requests OTP verification sent to parent's phone to authenticate access.\n3.\tParent inputs the OTP.\n4.\tSystem validates the OTP.\n5.\tSystem retrieves and displays the latest VISIT notes (Diagnosis, Doctor's advice) and PRESCRIPTION details (Medication name, dosage, frequency).\n6.\tParent reads the instructions.",
                "alt": "Step 4_OTP Invalid\n1.\tParent inputs an incorrect or expired OTP (MSG04).\n2.\tSystem displays an error message and asks to re-enter OTP or request a new one.\n3.\tReturn to step 3 of normal flow."
            }
        ]
    },
    {
        "feature": "2.4 Appointment Management",
        "ucs": [
            {
                "title": "2.4.1 UC-05 View Appointment Dashboard",
                "primary_actor": "Receptionist", "secondary_actor": "HCMS System",
                "desc": "As a receptionist, I want to see an overview of all appointments for the day, to effectively coordinate patient flow.",
                "pre": "Receptionist is securely logged into the system.",
                "post": "●\tSystem displays current day's APPOINTMENT list.",
                "norm": "View Dashboard\n1.\tReceptionist clicks on the 'Dashboard' tab from the sidebar.\n2.\tSystem queries the APPOINTMENT list for the current date.\n3.\tSystem categorizes and displays the queue ('Pending', 'Confirmed', 'Checked-in', 'Completed').\n4.\tReceptionist views the queue to monitor clinic traffic.",
                "alt": "None"
            },
            {
                "title": "2.4.2 UC-06 Confirm Appointment",
                "primary_actor": "Receptionist", "secondary_actor": "HCMS System",
                "desc": "As a receptionist, I want to confirm pending online bookings to ensure schedule accuracy and notify parents.",
                "pre": "Dashboard displays a 'Pending' APPOINTMENT.",
                "post": "●\tAPPOINTMENT status becomes 'Confirmed'.",
                "norm": "Confirm Booking\n1.\tReceptionist reviews the 'Pending' APPOINTMENT details on the Dashboard.\n2.\tReceptionist clicks the 'Confirm' button for that appointment.\n3.\tSystem prompts for action confirmation.\n4.\tReceptionist confirms the action.\n5.\tSystem updates the APPOINTMENT status from 'Pending' to 'Confirmed'.\n6.\tSystem automatically triggers UC-03 (Receive Appointment Confirmation).",
                "alt": "Step 2.1_Cancel Booking\n1.\tReceptionist identifies a duplicate or invalid booking.\n2.\tReceptionist clicks 'Cancel' and provides a reason.\n3.\tSystem updates the APPOINTMENT status to 'Cancelled'.\n4.\tReturn to step 6 of normal flow (triggers Cancellation notification)."
            },
            {
                "title": "2.4.3 UC-07 Register Walk-in Appointment",
                "primary_actor": "Receptionist", "secondary_actor": "HCMS System",
                "desc": "As a receptionist, I want to register walk-in patients quickly at the desk to maintain accurate queues.",
                "pre": "Parent arrives at the clinic without prior booking.",
                "post": "●\tAPPOINTMENT and PATIENT records are created.\n●\tAPPOINTMENT status is 'Checked-in'.",
                "norm": "Register Walk-in\n1.\tReceptionist clicks 'New Walk-in' button on the Dashboard.\n2.\tSystem displays the Walk-in Registration form.\n3.\tReceptionist inputs new patient demographics (Name, DOB, Parent Contact).\n4.\tReceptionist assigns the patient to an available Doctor queue.\n5.\tReceptionist clicks 'Register'.\n6.\tSystem saves a new PATIENT entity.\n7.\tSystem creates a new APPOINTMENT with status 'Checked-in'.\n8.\tSystem updates the Clinical Dashboard for the respective Doctor.",
                "alt": "Step 3.1_Patient Exists\n1.\tReceptionist selects the 'Search Existing' option and search by phone number.\n2.\tSystem displays existing patient details.\n3.\tReceptionist selects the patient to autofill demographics.\n4.\tReturn to step 4 of normal flow."
            }
        ]
    },
    {
        "feature": "2.5 Billing & Payment",
        "ucs": [
            {
                "title": "2.5.1 UC-08 View Billing Invoice",
                "primary_actor": "Receptionist", "secondary_actor": "HCMS System",
                "desc": "As a receptionist, I want to view the auto-generated checkout invoice for a patient to collect payment securely.",
                "pre": "Doctor finalized the VISIT, and System generated the BILLING record.",
                "post": "●\tInvoice is displayed accurately to the Receptionist.",
                "norm": "View Checkout\n1.\tReceptionist navigates to the 'Billing' tab or selects the 'Ready for Checkout' appointment on the Dashboard.\n2.\tSystem retrieves the corresponding BILLING and PRESCRIPTION data.\n3.\tSystem displays the detailed invoice including consultation fee and medication costs.\n4.\tReceptionist prints the invoice or communicates the total amount to the Parent.",
                "alt": "None"
            },
            {
                "title": "2.5.2 UC-09 Update Payment Status (1-Click Paid)",
                "primary_actor": "Receptionist", "secondary_actor": "Parents",
                "desc": "As a receptionist, I want to mark the bill as paid with 1-click to ensure fast checkout operations.",
                "pre": "Receptionist views the 'Unpaid' invoice (UC-08).",
                "post": "●\tBILLING status is updated to 'Paid'.\n●\tSystem logs checkout transaction.",
                "norm": "Process Payment\n1.\tParent provides cash or completes a bank transfer for the total invoice amount.\n2.\tReceptionist verifies the payment received.\n3.\tReceptionist clicks the 'Mark as Paid' (1-Click) button on the system.\n4.\tSystem updates the BILLING status from 'Unpaid' to 'Paid'.\n5.\tSystem records the transaction for daily revenue tracking.",
                "alt": "Step 1.1_Payment Delayed/Failed\n1.\tParent does not have enough funds or bank transfer fails.\n2.\tReceptionist leaves the invoice status as 'Unpaid'.\n3.\tSystem keeps the invoice in the 'Pending Payment' list."
            }
        ]
    },
    {
        "feature": "2.6 Clinical Dashboard",
        "ucs": [
            {
                "title": "2.6.1 UC-10 View Patient Queue",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to view the queue of waiting patients assigned to me so I know who to examine next.",
                "pre": "Doctor is logged in to the Clinical Dashboard.",
                "post": "●\tSystem displays the queue of patients.",
                "norm": "View Queue\n1.\tDoctor opens the Clinical Dashboard screen.\n2.\tSystem retrieves the daily APPOINTMENTs assigned to the currently logged in Doctor.\n3.\tSystem displays the list categorized by 'Checked-in' (Waiting) and 'Completed'.\n4.\tDoctor selects the next 'Checked-in' patient from the top of the queue to begin consultation.",
                "alt": "None"
            }
        ]
    },
    {
        "feature": "2.7 EMR Access",
        "ucs": [
            {
                "title": "2.7.1 UC-11 Access Patient Medical History",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to review past visits, diagnoses, and allergies to make informed clinical decisions.",
                "pre": "Doctor selects a patient from the queue.",
                "post": "●\tSystem displays the EMR Summary.",
                "norm": "View History\n1.\tDoctor clicks on the Patient's name in the Clinical Dashboard queue.\n2.\tSystem retrieves the patient's demographic profile, registered allergies, and past VISIT/PRESCRIPTION records.\n3.\tSystem displays the historical summary in a structured timeline view.\n4.\tDoctor reviews the information before proceeding to the current consultation notes.",
                "alt": "None"
            }
        ]
    },
    {
        "feature": "2.8 EMR \u2013 Visit & Prescription Management",
        "ucs": [
            {
                "title": "2.8.1 UC-12 Record Clinical Consultation",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to type clinical notes and diagnoses directly into the EMR to eliminate paper files.",
                "pre": "Doctor has reviewed the patient's history (UC-11).",
                "post": "●\tVISIT entity is created and securely linked to the APPOINTMENT.",
                "norm": "Enter Notes\n1.\tDoctor clicks 'Start Consultation'.\n2.\tSystem opens a new VISIT record form interface.\n3.\tDoctor inputs Symptoms, Vital Signs, and Diagnosis.\n4.\tDoctor clicks 'Save Consultation'.\n5.\tSystem validates the inputs (e.g., diagnosis cannot be empty).\n6.\tSystem securely stores the VISIT record into the database.",
                "alt": "Step 5_Validation Error\n1.\tDoctor forgets to input the mandatory Diagnosis field (MSG05).\n2.\tSystem highlights the missing required field and prevents saving.\n3.\tReturn to step 3 of normal flow."
            },
            {
                "title": "2.8.2 UC-13 Issue Electronic Prescription",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to quickly add medications using autocomplete to save time when prescribing.",
                "pre": "Doctor is in an active VISIT session.",
                "post": "●\tPRESCRIPTION entity is created.\n●\tBILLING entity is automatically generated.",
                "norm": "Create Prescription\n1.\tDoctor clicks 'Add Prescription'.\n2.\tDoctor types the first few letters of a medication name into the drug search field.\n3.\tSystem displays a dropdown with Autocomplete suggestions from the drug dictionary.\n4.\tDoctor selects the desired medication from the suggestions.\n5.\tDoctor inputs dosage instructions (e.g., '2 times/day').\n6.\tDoctor repeats steps 2-5 for other medications, then clicks 'Finalize Consultation'.\n7.\tSystem saves the PRESCRIPTION linked to the active VISIT.\n8.\tSystem automatically generates a BILLING invoice based on the consultation fee and medication costs.\n9.\tSystem updates APPOINTMENT status to 'Completed'.",
                "alt": "Step 6.1_No Medication Needed\n1.\tDoctor determines no medicine is needed for this patient.\n2.\tDoctor clicks 'Finalize Consultation' without adding any drugs.\n3.\tReturn to step 8 of normal flow (System generates a Billing invoice containing consultation fee only)."
            }
        ]
    },
    {
        "feature": "2.9 EMR \u2013 Patient Management",
        "ucs": [
            {
                "title": "2.9.1 UC-14 Manage Patient Profile",
                "primary_actor": "Doctor", "secondary_actor": "HCMS System",
                "desc": "As a doctor, I want to correct or update the core demographic and allergy information of the patient to ensure safety.",
                "pre": "Doctor views the Patient Medical History (UC-11).",
                "post": "●\tPATIENT entity is updated successfully.",
                "norm": "Edit Profile\n1.\tDoctor clicks 'Edit Profile' on the patient's information card.\n2.\tSystem makes the data fields (Height, Weight, Allergies) editable.\n3.\tDoctor types in the updated information (e.g., adding an allergy to Penicillin).\n4.\tDoctor clicks 'Update'.\n5.\tSystem saves the modified PATIENT record to the database.\n6.\tSystem displays a success banner confirming the update.",
                "alt": "None"
            }
        ]
    }
]

def format_cell(cell, text, bold=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.bold = bold

def create_table_element(uc):
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    format_cell(table.cell(0, 0), 'Primary Actors')
    format_cell(table.cell(0, 1), uc['primary_actor'])
    format_cell(table.cell(0, 2), 'Secondary Actors')
    format_cell(table.cell(0, 3), uc['secondary_actor'])
    
    format_cell(table.cell(1, 0), 'Description')
    table.cell(1, 1).merge(table.cell(1, 3))
    format_cell(table.cell(1, 1), uc['desc'])
    
    format_cell(table.cell(2, 0), 'Preconditions')
    table.cell(2, 1).merge(table.cell(2, 3))
    format_cell(table.cell(2, 1), uc['pre'])
    
    format_cell(table.cell(3, 0), 'Postconditions')
    table.cell(3, 1).merge(table.cell(3, 3))
    format_cell(table.cell(3, 1), uc['post'])
    
    format_cell(table.cell(4, 0), 'Normal Sequence/Flow')
    table.cell(4, 1).merge(table.cell(4, 3))
    format_cell(table.cell(4, 1), uc['norm'])
    
    format_cell(table.cell(5, 0), 'Alternative Sequences/Flows')
    table.cell(5, 1).merge(table.cell(5, 3))
    format_cell(table.cell(5, 1), uc['alt'])
    
    return table._tbl

body = doc._body._body

start_elem = None
end_elem = None

for p in doc.paragraphs:
    txt = p.text.strip()
    if txt.startswith("2.1 <<Feature Name1>>"):
        start_elem = p._element
    elif txt.startswith("2.2 Xyz Feature"):
        end_elem = p._element

if start_elem is not None and end_elem is not None:
    elements_to_delete = []
    in_range = False
    for elem in body:
        if elem == start_elem:
            in_range = True
        if elem == end_elem:
            in_range = False
        if in_range:
            elements_to_delete.append(elem)
            
    for elem in elements_to_delete:
        body.remove(elem)
        
    for feature_obj in ucs_data:
        feat_p = doc.add_paragraph(feature_obj["feature"])
        feat_p.style = doc.styles['Heading 2']
        for r in feat_p.runs: r.font.name = 'Arial'
        end_elem.addprevious(feat_p._p)
        
        for uc in feature_obj["ucs"]:
            uc_p = doc.add_paragraph(uc["title"])
            uc_p.style = doc.styles['Heading 3']
            for r in uc_p.runs: 
                r.font.name = 'Arial'
                r.bold = True
            end_elem.addprevious(uc_p._p)
            
            tbl_xml = create_table_element(uc)
            end_elem.addprevious(tbl_xml)
            
            spacer = doc.add_paragraph("")
            end_elem.addprevious(spacer._p)

    doc.save(file_path)
    print("Done! Extended and replaced 14 use cases with highly detailed specifications.")
else:
    print("Error: Could not find markers.")
