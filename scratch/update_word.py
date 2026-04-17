import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

file_path = r'e:\Documents\Ky_8\Co Dieu AISDLC\HCMS---Healthcare-Clinic-System\Document\02_Requirements\HCMS_SRS_v0.1.docx'
doc = docx.Document(file_path)

usecases = [
    {
        "title": "2.1.1 Self-Service Booking Portal",
        "primary_actor": "Parents",
        "secondary_actor": "HCMS System, Receptionist",
        "description": "As a parent, I want to book an appointment with a doctor online so that I can save time and avoid waiting to call the clinic.",
        "preconditions": "User accesses the Self-Service Portal.",
        "postconditions": "•\tUser books an appointment successfully\n•\tThe system creates an APPOINTMENT with status 'Pending'\n•\tThe system displays the appointment on the Receptionist's Dashboard",
        "normal_flow": "Book Appointment\n1.\tParent accesses the booking portal and chooses a preferred Doctor and available time slot.\n2.\tSystem shows the patient demographic and reason for visit form.\n3.\tParent types in the patient's information and reason for visit.\n4.\tParent clicks the 'Book' button.\n5.\tSystem validates the inputted information.\n6.\tSystem creates a new APPOINTMENT with status 'Pending'.\n7.\tSystem displays the pending appointment on the Receptionist's Dashboard.\n8.\tSystem redirects the parent to the Booking Confirmation page.",
        "alt_flow": "Step 5_Invalid Input Data\nIf the parent leaves mandatory fields blank or inputs invalid data:\n1.\tSystem displays a relevant error message highlighting the invalid fields.\n2.\tReturn to step 3 of the normal flow."
    },
    {
        "title": "2.1.2 Clinical Consultation & EMR",
        "primary_actor": "Doctor",
        "secondary_actor": "HCMS System",
        "description": "As a doctor, I want to view a patient's medical history, record clinical notes, and electronically prescribe medication so that I can efficiently manage patient care and reduce paperwork.",
        "preconditions": "The doctor is logged into the system and the patient has an active appointment.",
        "postconditions": "•\tThe patient's VISIT record is saved with clinical notes.\n•\tA PRESCRIPTION is generated and linked to the VISIT.",
        "normal_flow": "Record Consultation\n1.\tDoctor opens the Dashboard and selects a waiting patient's appointment.\n2.\tSystem displays the patient's Electronic Medical Record (EMR) including history.\n3.\tDoctor inputs symptoms, diagnosis, and clinical notes into the VISIT record.\n4.\tDoctor proceeds to add a prescription and starts typing medication names.\n5.\tSystem provides Autocomplete suggestions for medication names.\n6.\tDoctor selects medications and finalizes the prescription.\n7.\tSystem securely saves the VISIT and PRESCRIPTION data.\n8.\tSystem allows the parent to view the consultation notes/prescription post-visit.",
        "alt_flow": "Step 4_No Prescription Required\n1.\tDoctor determines that no medication is needed for the patient.\n2.\tDoctor finalizes the consultation without adding any items to the prescription.\n3.\tReturn to step 7 of the normal flow."
    },
    {
        "title": "2.1.3 Streamlined 1-Click Checkout",
        "primary_actor": "Receptionist",
        "secondary_actor": "HCMS System, Parents",
        "description": "As a receptionist, I want to quickly process payments for completed visits using a 1-click checkout so that I can ensure fast billing without complex insurance handling.",
        "preconditions": "The doctor has finalized the patient's VISIT and PRESCRIPTION.",
        "postconditions": "•\tThe BILLING status is updated from Unpaid to Paid.\n•\tRevenue is recorded in the system for the day's financials.",
        "normal_flow": "Process Checkout\n1.\tSystem automatically generates a BILLING record based on the service fees and prescribed medications once the doctor finishes the visit.\n2.\tReceptionist views the automatically generated invoice on the Dashboard.\n3.\tReceptionist confirms the total amount with the Parent at the clinic counter.\n4.\tParent completes the payment.\n5.\tReceptionist clicks the 'Paid' (1-Click) button on the invoice.\n6.\tSystem updates the BILLING status from 'Unpaid' to 'Paid'.\n7.\tSystem logs the transaction for daily revenue tracking.",
        "alt_flow": "Step 4_Payment Delays\n1.\tIf the parent fails to or delays the payment, Receptionist leaves the invoice as 'Unpaid'.\n2.\tSystem retains the unpaid invoice on the Dashboard for later processing."
    }
]

def format_cell(cell, text, bold=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.bold = bold

def insert_table(uc, after_element):
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    format_cell(table.cell(0, 0), 'Primary Actors')
    format_cell(table.cell(0, 1), uc['primary_actor'])
    format_cell(table.cell(0, 2), 'Secondary Actors')
    format_cell(table.cell(0, 3), uc['secondary_actor'])
    
    format_cell(table.cell(1, 0), 'Description')
    table.cell(1, 1).merge(table.cell(1, 3))
    format_cell(table.cell(1, 1), uc['description'])
    
    format_cell(table.cell(2, 0), 'Preconditions')
    table.cell(2, 1).merge(table.cell(2, 3))
    format_cell(table.cell(2, 1), uc['preconditions'])
    
    format_cell(table.cell(3, 0), 'Postconditions')
    table.cell(3, 1).merge(table.cell(3, 3))
    format_cell(table.cell(3, 1), uc['postconditions'])
    
    format_cell(table.cell(4, 0), 'Normal Sequence/Flow')
    table.cell(4, 1).merge(table.cell(4, 3))
    format_cell(table.cell(4, 1), uc['normal_flow'])
    
    format_cell(table.cell(5, 0), 'Alternative Sequences/Flows')
    table.cell(5, 1).merge(table.cell(5, 3))
    format_cell(table.cell(5, 1), uc['alt_flow'])
    
    after_element.addnext(table._tbl)
    return table

# find the target paragraph
target_para = None
for p in doc.paragraphs:
    if "2.1 <<Feature Name1>>" in p.text:
        target_para = p
        p.text = "2.1 Operational Workflows"
        # Optional: set heading style
        p.style = doc.styles['Heading 2']
        break

if target_para:
    # First, let's remove the placeholder lines below target_para until '2.1.3 UC Name2'
    # Actually, modifying docx while iterating is dangerous. We will just delete specific ones.
    # To keep it simple, we don't delete them right away, but we just insert our tables in order,
    # and then we'll clean up the known old titles like "2.1.1 UC Name1", "2.1.2 Login System", "2.1.3 UC Name2", etc.
    
    # Let's clean up specifically the ones from the template
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt in ["2.1.1 UC Name1", "2.1.2 Login System", "2.1.3 UC Name2"] or \
           txt.startswith("Primary and Secondary Actors") or \
           txt.startswith("An actor is a person") or \
           txt.startswith("Description") or \
           txt.startswith("Provide a brief description") or \
           txt.startswith("Preconditions") or \
           txt.startswith("List any activities") or \
           txt.startswith("Postconditions") or \
           txt.startswith("Describe the state") or \
           txt.startswith("Normal Flow") or \
           txt.startswith("Provide a description") or \
           txt.startswith("Alternative Flows") or \
           txt.startswith("Describe below two") or \
           txt.startswith("Other successful usage") or \
           txt.startswith("Any anticipated error"):
            # Instead of deleting (which can mess up XML), clear the text
            # Wait, clearing text leaves empty lines. Let's delete the XML element for safety.
            p._element.getparent().remove(p._element)
    
    current_element = target_para._p
    
    for uc in usecases:
        # Create heading Title for UC
        new_p = doc.add_paragraph(uc['title'])
        new_p.style = doc.styles['Heading 3'] # usually 2.1.1 is Heading 3
        # Add bold to it
        for run in new_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.bold = True
            
        current_element.addnext(new_p._p)
        
        # Insert table
        new_table = insert_table(uc, new_p._p)
        current_element = new_table._tbl
        
        # Insert an empty line for spacing
        space_p = doc.add_paragraph("")
        current_element.addnext(space_p._p)
        current_element = space_p._p

doc.save(file_path)
print("Updated successfully")
