import docx

doc = docx.Document("Document/02_Requirements/HCMS_SRS_v0.1.docx")

for i, p in enumerate(doc.paragraphs):
    if "3.1 Feature Name1" in p.text:
        print(f"Found '3.1 Feature Name1' at index {i}")
    if "3.2 User Authentication" in p.text:
        print(f"Found '3.2 User Authentication' at index {i}")
    if "3.3 System Administration" in p.text:
        print(f"Found '3.3 System Administration' at index {i}")

for i, t in enumerate(doc.tables):
    try:
        cell_text = t.cell(0, 0).text
        if "Field Name" in cell_text:
            print(f"Found potential template table at index {i} with first cell: {cell_text}")
    except:
        pass
