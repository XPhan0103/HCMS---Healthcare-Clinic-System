import docx
import copy
from docx.table import Table

doc = docx.Document("Document/02_Requirements/HCMS_SRS_v0.1.docx")
p_target = doc.paragraphs[120] # "3.2 User Authentication"

t_template = doc.tables[20]

new_tbl = copy.deepcopy(t_template._tbl)
p_target._p.addprevious(new_tbl)
table_obj = Table(new_tbl, p_target._parent)
table_obj.cell(1, 0).text = "Test Field"
table_obj.cell(1, 1).text = "Test Description"

doc.save("scratch_test_table.docx")
