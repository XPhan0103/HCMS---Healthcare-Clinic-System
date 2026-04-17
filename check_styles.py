import docx

doc = docx.Document("Document/02_Requirements/HCMS_SRS_v0.1.docx")

p_3_1 = doc.paragraphs[109]
p_3_1_1 = doc.paragraphs[110]
p_3_1_1_1 = doc.paragraphs[111]

print("3.1 style:", p_3_1.style.name)
print("3.1.1 style:", p_3_1_1.style.name)
print("3.1.1.1 style:", p_3_1_1_1.style.name)
